import pdfplumber
from docx import Document as DocxDocument
import openpyxl
import pandas as pd
from PIL import Image
import os


def parse_file(file_path: str, file_type: str) -> str:
    """Route by file_type to the right extractor. Always returns a string."""
    parsers = {
        "pdf": parse_pdf,
        "xlsx": parse_excel, "xls": parse_excel,
        "csv": parse_csv,
        "docx": parse_docx, "doc": parse_docx,
        "txt": parse_text,
        "png": parse_image, "jpg": parse_image, "jpeg": parse_image,
    }
    parser = parsers.get(file_type.lower(), parse_text)
    try:
        return parser(file_path)
    except Exception as e:
        return f"[Error parsing {file_type} file: {str(e)}]"


def parse_pdf(file_path: str) -> str:
    """
    Use pdfplumber. For each page:
    1. Extract text with layout preservation
    2. Extract tables → convert each to a pipe-delimited table string
    Separate pages with "--- PAGE {n} ---"
    """
    pages = []
    with pdfplumber.open(file_path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables()
            table_text = ""
            for table in tables:
                rows = []
                for row in table:
                    rows.append(" | ".join(str(cell or "") for cell in row))
                table_text += "\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]\n"
            pages.append(f"--- PAGE {i+1} ---\n{text}\n{table_text}")
    return "\n\n".join(pages)


def parse_excel(file_path: str) -> str:
    """
    Use openpyxl. For each sheet:
    1. Read all rows into a list of lists
    2. Detect type by keywords in first 5 rows
    3. Format as: "=== SHEET: {name} (Detected: {type}) ===\n" + rows as pipe-delimited
    """
    wb = openpyxl.load_workbook(file_path, data_only=True)
    sheets_text = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            rows.append(" | ".join(str(cell if cell is not None else "") for cell in row))
        text = "\n".join(rows)
        # Detect sheet type
        sample = text[:500].lower()
        if any(kw in sample for kw in ["revenue", "sales", "net income", "gross profit"]):
            detected = "Income Statement"
        elif any(kw in sample for kw in ["assets", "liabilities", "equity"]):
            detected = "Balance Sheet"
        elif any(kw in sample for kw in ["cash flow", "operating", "investing"]):
            detected = "Cash Flow Statement"
        else:
            detected = "Financial Data"
        sheets_text.append(f"=== SHEET: {sheet_name} (Detected: {detected}) ===\n{text}")
    return "\n\n".join(sheets_text)


import csv

def parse_csv(file_path: str) -> str:
    """Read CSV using python's built-in csv module to handle variable columns, title rows, and multi-headers gracefully."""
    try:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                content = f.read()
                
        import io
        reader = csv.reader(io.StringIO(content))
        rows = []
        for row in reader:
            if any(cell.strip() for cell in row):
                # Only include non-empty cells or just join all cells properly
                rows.append(" | ".join(str(cell).strip() for cell in row))
        return "\n".join(rows)
    except Exception:
        return parse_text(file_path)


def parse_docx(file_path: str) -> str:
    """Extract paragraphs and tables from Word documents."""
    doc = DocxDocument(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text for cell in row.cells))
        parts.append("[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
    return "\n\n".join(parts)


def parse_text(file_path: str) -> str:
    """Read as plain text, try UTF-8 then latin-1."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="latin-1") as f:
            return f.read()


def parse_image(file_path: str) -> str:
    """Return placeholder — images handled by Claude vision during extraction."""
    img = Image.open(file_path)
    return f"[IMAGE: {os.path.basename(file_path)} — {img.width}x{img.height}px. Send to Claude Vision for extraction.]"
